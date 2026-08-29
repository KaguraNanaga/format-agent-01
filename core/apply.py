# 执行器接线（PLAN.md 第 7 节）：
# apply_format(docx_path, spec, rolemap, out_path) -> changelog
# 对每个段落按 RoleMap 取角色、从 FormatSpec 取规则，调用 core/executor.py 的
# 确定性函数改 XML；页边距/行网格走 section 级别。LLM 不碰 docx。

from docx import Document
from docx.shared import Mm

from core.executor import (
    set_doc_grid,
)
from core.style_set import (
    apply_named_style,
    clear_invalid_numbering_override,
    ensure_role_styles,
    resolve_target_body_style,
)


def apply_format(docx_path, spec, rolemap, out_path):
    """应用 FormatSpec × RoleMap，输出 docx，返回 changelog list[dict]。
    rolemap: {idx: role}（idx 对应 extract.py 的段落序号）。
    模板未明确指定的角色统一绑定目标文档原有正文样式，不借用 other 规则。
    表格内段落（idx >= len(doc.paragraphs)）v1 跳过。
    """
    doc = Document(docx_path)
    roles = spec.get("roles", {})
    # 必须在创建/更新 FormatAgent 样式之前解析，避免把新样式误认成目标原样式。
    target_body_style = resolve_target_body_style(doc, rolemap)
    role_styles = ensure_role_styles(
        doc, spec, target_body_style=target_body_style)

    # ---- 页面级 ----
    page = spec.get("page") or {}
    margin = page.get("margin") or {}
    section = doc.sections[0]
    if margin.get("top_mm") is not None:
        section.top_margin = Mm(margin["top_mm"])
    if margin.get("bottom_mm") is not None:
        section.bottom_margin = Mm(margin["bottom_mm"])
    if margin.get("left_mm") is not None:
        section.left_margin = Mm(margin["left_mm"])
    if margin.get("right_mm") is not None:
        section.right_margin = Mm(margin["right_mm"])
    line_grid = page.get("line_grid") or {}
    if line_grid.get("line_pt") is not None:
        set_doc_grid(doc, line_pt=line_grid["line_pt"])

    # ---- 段落级 ----
    changelog = []
    for idx, p in enumerate(doc.paragraphs):
        role = rolemap.get(idx, rolemap.get(str(idx)))
        if role is None:
            continue  # 未被标注的段落不动
        if role in roles:
            rule = roles[role]
            style = role_styles[role]
            changed = apply_named_style(p, style, rule, role=role)
            fallback_to_target_body = False
        else:
            # 只替换 pStyle：真实自动编号、段落直接格式和 run 内强调均保留；
            # 仅清掉 numId=0/ilvl=-1 这类会遮蔽正文缩进的“取消编号”残留。
            # 即使模板定义了 other，也不能把缺失的 date/signature 等角色借给它。
            invalid_numbering_removed = clear_invalid_numbering_override(p)
            p.style = target_body_style
            style = target_body_style
            changed = ["paragraph_style", "fallback_to_target_body"]
            if invalid_numbering_removed:
                changed.append("invalid_numbering_removed")
            fallback_to_target_body = True
        changelog.append({
            "idx": idx,
            "role": role,
            "style_name": style.name,
            "text": p.text.strip()[:30],
            "changed_fields": changed,
            "fallback_to_target_body": fallback_to_target_body,
        })

    doc.save(out_path)
    return changelog


def write_report(changelog, spec, report_path):
    """把 changelog 写成 markdown 修改对照报告。"""
    lines = ["# 排版修改对照报告", ""]
    page = spec.get("page") or {}
    if page:
        lines.append("## 页面设置")
        margin = page.get("margin") or {}
        if margin:
            lines.append(
                f"- 页边距（mm）：上 {margin.get('top_mm', '-')} / 下 {margin.get('bottom_mm', '-')}"
                f" / 左 {margin.get('left_mm', '-')} / 右 {margin.get('right_mm', '-')}")
        lg = page.get("line_grid") or {}
        if lg.get("line_pt"):
            lines.append(f"- 行网格：{lg['line_pt']} 磅/行")
        lines.append("")
    lines.append("## 段落修改明细")
    lines.append("")
    lines.append("| 段落 | 角色 | Word 样式 | 改动字段 | 内容摘要 |")
    lines.append("|---|---|---|---|---|")
    for c in changelog:
        fields = ", ".join(c["changed_fields"]) if c["changed_fields"] else "（无字段改动）"
        lines.append(
            f"| {c['idx']} | {c['role']} | {c.get('style_name', '-')} | {fields} | {c['text']} |")
    lines.append("")
    with open(report_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return report_path
