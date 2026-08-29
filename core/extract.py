# docx → 段落清单（PLAN.md 第 7 节）。
# 每段输出 {idx, text(截前80字), size_pt, bold, alignment, style_name, in_table}。
# size_pt/bold 用 effective_props 读"生效属性"而非样式名。
# in_table=True 的段落 v1 不参与重排，但仍列出（供角色标注参考上下文）。

from docx import Document

from core.effective_props import get_paragraph_effective_font

_ALIGN_MAP = {0: "left", 1: "center", 2: "right", 3: "justify"}


def _alignment_name(paragraph):
    a = paragraph.alignment
    if a is None:
        return None
    return _ALIGN_MAP.get(int(a), str(a))


def _spacing_pt(paragraph):
    """读段落直刷的段前/段后距（磅）。样式层定义的间距读不到，返回 None。"""
    pf = paragraph.paragraph_format
    sb = round(pf.space_before.pt, 1) if pf.space_before is not None else None
    sa = round(pf.space_after.pt, 1) if pf.space_after is not None else None
    return sb, sa


def _para_record(idx, p, in_table):
    eastasia, size_pt, bold = get_paragraph_effective_font(p)
    sb, sa = _spacing_pt(p)
    return {
        "idx": idx,
        "text": p.text.strip()[:80],
        "size_pt": size_pt,
        "bold": bold,
        "alignment": _alignment_name(p),
        "style_name": p.style.name if p.style is not None else None,
        "space_before_pt": sb,
        "space_after_pt": sa,
        "in_table": in_table,
    }


def extract_paragraphs(docx_path):
    """返回段落清单 list[dict]，idx 为全文段落序号（含表格内段落）。"""
    doc = Document(docx_path)
    out = []
    for idx, p in enumerate(doc.paragraphs):
        out.append(_para_record(idx, p, False))
    # 表格段落追加在末尾（doc.paragraphs 不含表格内段落，单独列出）
    idx = len(out)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    out.append(_para_record(idx, p, True))
                    idx += 1
    return out


if __name__ == "__main__":
    import json
    import sys
    paras = extract_paragraphs(sys.argv[1])
    print(json.dumps(paras, ensure_ascii=False, indent=2))
