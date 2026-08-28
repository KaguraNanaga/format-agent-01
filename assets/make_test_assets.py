# 生成测试素材：一份"格式全错"的演示文档 messy.docx + 规范文字 spec.txt + 标准答案 spec_std.json。
# 用法: python assets/make_test_assets.py
import json
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))

PARAS = [
    # (text, size_pt, bold, alignment) —— 故意全错：宋体五号、无缩进、标题靠左
    ("关于开展2026年度数字化转型工作专项检查的通知", 12, False, WD_ALIGN_PARAGRAPH.LEFT),
    ("各有关单位：", 12, False, WD_ALIGN_PARAGRAPH.LEFT),
    ("一、检查目的", 12, True, WD_ALIGN_PARAGRAPH.LEFT),
    ("为深入贯彻落实数字化转型战略部署，全面掌握各单位数字化建设工作进展情况，"
     "及时发现和解决工作中存在的突出问题，经研究，决定开展2026年度数字化转型工作专项检查。", 12, False, WD_ALIGN_PARAGRAPH.LEFT),
    ("二、检查内容", 12, True, WD_ALIGN_PARAGRAPH.LEFT),
    ("（一）组织领导情况。重点检查各单位数字化转型工作领导小组组建及运行情况。", 12, False, WD_ALIGN_PARAGRAPH.LEFT),
    ("（二）项目推进情况。重点检查年度重点项目立项、实施、验收各环节管理情况。", 12, False, WD_ALIGN_PARAGRAPH.LEFT),
    ("请各单位高度重视，认真组织开展自查，并于2026年9月15日前将自查报告报送至市数字化局。", 12, False, WD_ALIGN_PARAGRAPH.LEFT),
    ("市数字化局办公室", 12, False, WD_ALIGN_PARAGRAPH.LEFT),
    ("2026年8月28日", 12, False, WD_ALIGN_PARAGRAPH.LEFT),
]

SPEC_TEXT = """公文排版规范：
1. 页面：A4 纸，上边距 37 毫米，下边距 35 毫米，左边距 28 毫米，右边距 26 毫米，行网格每行 28 磅。
2. 标题：方正小标宋简体，二号（22 磅），不加粗，居中，行距固定值 28 磅。
3. 一级标题：黑体，三号（16 磅），左对齐，首行缩进 2 字符，行距固定值 28 磅。
4. 正文：仿宋_GB2312，西文 Times New Roman，三号（16 磅），两端对齐，首行缩进 2 字符，行距固定值 28 磅。
5. 落款（署名）：仿宋_GB2312，三号（16 磅），右对齐，行距固定值 28 磅。
6. 成文日期：仿宋_GB2312，三号（16 磅），右对齐，行距固定值 28 磅。
"""

# 标准答案 FormatSpec（验收基准，也是 LLM 失败时的降级输入）
SPEC_STD = {
    "page": {
        "size": "A4",
        "margin": {"top_mm": 37, "bottom_mm": 35, "left_mm": 28, "right_mm": 26},
        "line_grid": {"line_pt": 28},
    },
    "roles": {
        "title": {"font_eastasia": "方正小标宋简体", "font_ascii": "Times New Roman",
                  "size_pt": 22, "bold": False, "alignment": "center",
                  "line_spacing": {"type": "exact", "pt": 28}},
        "heading_1": {"font_eastasia": "黑体", "size_pt": 16, "alignment": "left",
                      "first_line_indent_chars": 2,
                      "line_spacing": {"type": "exact", "pt": 28}},
        "body": {"font_eastasia": "仿宋_GB2312", "font_ascii": "Times New Roman",
                 "size_pt": 16, "alignment": "justify",
                 "first_line_indent_chars": 2,
                 "line_spacing": {"type": "exact", "pt": 28}},
        "signature": {"font_eastasia": "仿宋_GB2312", "size_pt": 16, "alignment": "right",
                      "line_spacing": {"type": "exact", "pt": 28}},
        "date": {"font_eastasia": "仿宋_GB2312", "size_pt": 16, "alignment": "right",
                 "line_spacing": {"type": "exact", "pt": 28}},
    },
}

# 标准答案 RoleMap（对照 PARAS 的顺序）
ROLEMAP_STD = {0: "title", 1: "body", 2: "heading_1", 3: "body", 4: "heading_1",
               5: "body", 6: "body", 7: "body", 8: "signature", 9: "date"}


def main():
    doc = Document()
    for text, size, bold, align in PARAS:
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = "宋体"
    docx_path = os.path.join(HERE, "messy.docx")
    doc.save(docx_path)

    with open(os.path.join(HERE, "spec.txt"), "w", encoding="utf-8") as f:
        f.write(SPEC_TEXT)
    with open(os.path.join(HERE, "spec_std.json"), "w", encoding="utf-8") as f:
        json.dump(SPEC_STD, f, ensure_ascii=False, indent=2)
    with open(os.path.join(HERE, "rolemap_std.json"), "w", encoding="utf-8") as f:
        json.dump({str(k): v for k, v in ROLEMAP_STD.items()}, f, ensure_ascii=False, indent=2)
    print("生成完毕: messy.docx / spec.txt / spec_std.json / rolemap_std.json")


if __name__ == "__main__":
    main()
