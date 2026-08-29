# 生成演示案例的"改之前"文档：格式故意做乱（宋体小字、标题靠左、无缩进、单倍行距）。
# 案例4 用真实文件（星驰议案），不在此生成。
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))


def make_messy(path, paragraphs):
    """paragraphs: [(text, kind)]，kind: title/subtitle/body/bold —— 全部用混乱格式写入。"""
    doc = Document()
    for text, kind in paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 标题也不居中，全部靠左
        pf = p.paragraph_format
        pf.line_spacing = 1.0  # 单倍行距
        pf.first_line_indent = None  # 无首行缩进
        run = p.add_run(text)
        if kind == "title":
            run.font.name = "微软雅黑"
            run.font.size = Pt(14)
            run.font.bold = True
        elif kind == "bold":
            run.font.name = "微软雅黑"
            run.font.size = Pt(11)
            run.font.bold = True
        else:
            run.font.name = "宋体"
            run.font.size = Pt(12)
    doc.save(path)
    return path


# ---------- 案例 1：规范文字 × 纯文本无结构 ----------
CASE1 = [
    ("关于组织开展2026年度职工健康体检的通知", "title"),
    ("各部门、各单位：", "body"),
    ("为保障广大职工身体健康，及时了解职工身体状况，做到疾病早发现、早预防、早治疗，"
     "经公司研究，决定于2026年10月中旬组织全体职工进行年度健康体检。", "body"),
    ("本次体检由公司统一联系体检机构承检，费用由公司承担。体检项目包括内科、外科、"
     "血常规、尿常规、肝功能、心电图、胸部X光等常规项目。", "body"),
    ("请各部门、各单位高度重视，合理安排工作，组织职工按时参加体检，并于2026年9月20日前"
     "将参加体检人员名单报送至综合办公室。", "body"),
    ("体检当天请空腹前往，携带本人身份证。体检报告由综合办公室统一发放，请妥善保管。", "body"),
    ("特此通知。", "body"),
    ("综合办公室", "body"),
    ("2026年9月1日", "body"),
]

# ---------- 案例 2：规范文字 × 有序号结构（手工序号） ----------
CASE2 = [
    ("2026年秋季安全生产大检查工作方案", "title"),
    ("为深入排查治理安全隐患，坚决防范和遏制各类生产安全事故发生，"
     "结合公司实际，制定本方案。", "body"),
    ("一、工作目标", "bold"),
    ("通过开展秋季安全生产大检查，全面排查各生产经营场所安全隐患，"
     "实现隐患整改闭环率百分之百，确保公司安全生产形势持续稳定。", "body"),
    ("二、检查范围与内容", "bold"),
    ("（一）检查范围", "bold"),
    ("公司所属各办公楼宇、仓储场所、在建工程项目及出租经营场所。", "body"),
    ("（二）检查内容", "bold"),
    ("1.设备设施安全", "body"),
    ("重点检查特种设备定期检验、日常维护保养及操作人员持证上岗情况。", "body"),
    ("2.消防安全", "body"),
    ("重点检查消防通道畅通、灭火器材配备有效及应急预案演练开展情况。", "body"),
    ("三、时间安排", "bold"),
    ("本次大检查自2026年9月10日开始，至9月30日结束，分自查自纠、"
     "集中检查、总结整改三个阶段进行。", "body"),
    ("四、工作要求", "bold"),
    ("各部门主要负责人要亲自部署、亲自督办，对排查出的隐患建立台账、"
     "限期整改，并于10月10日前将整改情况书面报送安全环保部。", "body"),
    ("安全环保部", "body"),
    ("2026年9月5日", "body"),
]

# ---------- 案例 3：党委会议题模板 × 纯文本无结构 ----------
CASE3 = [
    ("关于建立公司数字化转型专家库的汇报", "title"),
    ("（议题提交人：王丽）", "body"),
    ("为推进公司数字化转型工作，充分发挥专家在技术评审、方案论证和项目验收中的"
     "支撑作用，拟建立公司数字化转型专家库。", "body"),
    ("专家库拟由内部专家和外部专家组成，内部专家从各部门技术骨干中遴选，"
     "外部专家从高校、科研院所和行业龙头企业中聘任，首批规模控制在三十人以内。", "body"),
    ("专家库实行动态管理，每两年评估一次，根据项目参与度和专业表现进行调整。"
     "专家参与评审工作按照公司相关规定支付劳务报酬。", "body"),
    ("本事项已征求人力资源部和财务部意见，无分歧。", "body"),
    ("拟提交总经理办公会审议。", "body"),
]


def main():
    made = []
    made.append(make_messy(os.path.join(HERE, "case1", "改前-健康体检通知.docx"), CASE1))
    made.append(make_messy(os.path.join(HERE, "case2", "改前-安全生产大检查方案.docx"), CASE2))
    made.append(make_messy(os.path.join(HERE, "case3", "改前-专家库汇报.docx"), CASE3))
    for p in made:
        print("生成:", p)


if __name__ == "__main__":
    main()
