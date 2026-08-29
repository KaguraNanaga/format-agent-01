# 生成四象限演示案例 v2 的输入素材（文体更丰富：新闻稿/周报/论文/党委议案）。
# case4（党委议案）沿用已有素材，此处只生成 case1~case3 的"改前"文档与模板。
import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))


def make_messy(path, paragraphs):
    """格式故意做乱：字体杂、标题靠左、无缩进、单倍行距。"""
    doc = Document()
    for text, kind in paragraphs:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(text)
        if kind == "title":
            run.font.name = "微软雅黑"
            run.font.size = Pt(14)
            run.font.bold = True
        elif kind == "bold":
            run.font.name = "微软雅黑"
            run.font.size = Pt(11)
            run.font.bold = True
        else:
            run.font.name = "宋体"
            run.font.size = Pt(12)
    doc.save(path)
    return path


def make_template(path, blocks):
    """格式正确的模板 docx（供模板读规则）。blocks: [(text, font, size, bold, align, spacing)]"""
    doc = Document()
    for text, font, size, bold, align, spacing in blocks:
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.line_spacing = spacing
        run = p.add_run(text)
        run.font.name = font
        run.font.size = Pt(size)
        run.font.bold = bold
        # 中文字体必须写 eastAsia，否则模板读不到
        rpr = run._element.get_or_add_rPr()
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.append(rfonts)
        rfonts.set(qn("w:eastAsia"), font)
    doc.save(path)
    return path


# ---------- 案例 1：新闻稿（无结构）× 自然语言规范 ----------
CASE1_SPEC = """新闻稿排版规范：
1. 页面：A4 纸，上边距 25.4 毫米，下边距 25.4 毫米，左边距 31.8 毫米，右边距 31.8 毫米。
2. 标题：黑体，二号（22 磅），加粗，居中。
3. 正文：宋体，西文 Times New Roman，小四（12 磅），两端对齐，首行缩进 2 字符，行距 1.5 倍。
4. 落款（发布单位）：宋体，小四（12 磅），右对齐。
5. 成文日期：宋体，小四（12 磅），右对齐。
"""

CASE1 = [
    ("智汇云平台2.0正式发布，中小企业数字化转型再添新工具", "title"),
    ("本报讯  9月10日，智汇信息技术有限公司在杭州正式发布智汇云平台2.0版本，"
     "面向中小企业提供一站式数字化转型解决方案。", "body"),
    ("新版平台在原有基础上全面升级了数据分析引擎，新增智能报表、供应链协同、"
     "移动审批三大模块，可帮助企业将日常运营数据实时转化为经营决策依据。", "body"),
    ("发布会上，公司首席技术官介绍，智汇云平台2.0采用轻量化部署模式，"
     "企业无需购置服务器，注册即可使用，首年使用成本较传统方案下降约六成。", "body"),
    ("据悉，该平台自2024年上线以来已服务企业用户超过两万家，覆盖制造、零售、"
     "物流等多个行业。本次2.0版本即日起开放申请，前一千家签约企业可享受三个月免费试用期。", "body"),
    ("智汇信息技术有限公司品牌市场部", "body"),
    ("2026年9月10日", "body"),
]

# ---------- 案例 2：项目周报（无结构）× Word 模板 ----------
CASE2 = [
    ("研发一组项目周报（9月第2周）", "title"),
    ("汇报人：陈晓", "body"),
    ("本周智能客服项目完成意图识别模型第二轮训练，准确率提升至百分之九十一，"
     "较上周提高三个百分点。", "body"),
    ("知识库管理系统完成前端页面开发，已进入联调阶段，目前联调进度约百分之六十。", "body"),
    ("本周遇到的主要问题是训练数据标注人力不足，已协调外包团队增援，预计下周缓解。", "body"),
    ("下周计划完成模型第三轮优化，启动知识库系统用户验收测试准备工作。", "body"),
    ("2026年9月12日", "body"),
]

# ---------- 案例 3：课程论文（有结构）× 自然语言规范 ----------
CASE3 = [
    ("数字经济对传统零售业的影响研究", "title"),
    ("摘要：随着数字技术的迅猛发展，数字经济正在深刻改变传统零售业的发展格局。"
     "本文从销售渠道、供应链、消费者行为三个维度分析数字经济对传统零售业的影响，"
     "并提出相应的转型建议。", "body"),
    ("关键词：数字经济；传统零售业；转型升级", "body"),
    ("一、引言", "bold"),
    ("（一）研究背景", "bold"),
    ("近年来，我国数字经济规模持续扩大，占国内生产总值比重不断提升。"
     "零售业作为连接生产与消费的重要环节，首当其冲地受到数字技术的冲击与重塑。", "body"),
    ("（二）研究意义", "bold"),
    ("研究数字经济对传统零售业的影响，有助于传统零售企业认清形势、把握机遇，"
     "对推动零售业高质量发展具有现实意义。", "body"),
    ("二、数字经济概述", "bold"),
    ("数字经济是以数据资源为关键生产要素、以现代信息网络为重要载体、"
     "以信息通信技术融合应用为推动力的经济形态。", "body"),
    ("三、数字经济对传统零售业的影响", "bold"),
    ("（一）销售渠道的变革", "bold"),
    ("1.线上消费规模持续扩大", "body"),
    ("电子商务平台打破了传统零售业的地域限制，消费者可以随时随地完成购物，"
     "线上零售额占社会消费品零售总额的比重逐年上升。", "body"),
    ("2.全渠道融合成为趋势", "body"),
    ("线上线下融合的运营模式成为主流，传统门店逐步转型为体验中心和前置仓。", "body"),
    ("（二）供应链的重构", "bold"),
    ("数字技术使供应链各环节信息实时共享，库存周转效率显著提升，"
     "以销定产的柔性供应模式逐渐取代传统的以产定销模式。", "body"),
    ("四、结论与建议", "bold"),
    ("传统零售企业应主动拥抱数字技术，加快数字化基础设施建设，"
     "重构以消费者为中心的运营体系，实现线上线下一体化发展。", "body"),
    ("参考文献", "bold"),
    ("[1] 陈明. 数字经济导论[M]. 北京: 经济科学出版社, 2024.", "body"),
    ("[2] 李华. 零售业数字化转型路径研究[J]. 商业经济研究, 2025(3): 45-52.", "body"),
]


def main():
    made = []
    with open(os.path.join(HERE, "case1", "规范文字-新闻稿排版规范.txt"), "w", encoding="utf-8") as f:
        f.write(CASE1_SPEC)
    made.append(make_messy(os.path.join(HERE, "case1", "改前-产品发布新闻稿.docx"), CASE1))

    # 案例2 模板：格式正确的周报模板（微软雅黑体系）
    made.append(make_template(
        os.path.join(HERE, "case2", "模板-项目周报.docx"),
        [
            ("项目周报标题（微软雅黑 16磅 加粗 居中）", "微软雅黑", 16, True, WD_ALIGN_PARAGRAPH.CENTER, 1.5),
            ("（汇报人：XXX）（微软雅黑 11磅 居中）", "微软雅黑", 11, False, WD_ALIGN_PARAGRAPH.CENTER, 1.5),
            ("正文段落示例。本周完成了哪些工作、遇到什么问题、下周计划如何，"
             "按自然段书写即可。（微软雅黑 11磅 左对齐 1.5倍行距）",
             "微软雅黑", 11, False, WD_ALIGN_PARAGRAPH.LEFT, 1.5),
            ("日期行示例：2026年X月X日（微软雅黑 11磅 右对齐）", "微软雅黑", 11, False, WD_ALIGN_PARAGRAPH.RIGHT, 1.5),
        ]))
    made.append(make_messy(os.path.join(HERE, "case2", "改前-研发周报.docx"), CASE2))

    made.append(make_messy(os.path.join(HERE, "case3", "改前-数字经济课程论文.docx"), CASE3))
    for p in made:
        print("生成:", p)


if __name__ == "__main__":
    main()
